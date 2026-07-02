"""DOCX text extraction using python-docx."""

import io

from docx import Document as DocxDocument

from app.core.exceptions import ValidationError
from app.services.parsers.base import DocumentParser


class DocxParser(DocumentParser):
    """Extracts text from Word (.docx) files, including tables."""

    def extract_text(self, data: bytes) -> str:
        """Extract text from all paragraphs and tables in a .docx file.

        Args:
            data: The raw .docx file bytes.

        Returns:
            Extracted text, with paragraphs and table rows separated by
            newlines.

        Raises:
            ValidationError: If the file cannot be parsed as a valid .docx.
        """
        try:
            document = DocxDocument(io.BytesIO(data))
        except Exception as exc:
            raise ValidationError("File could not be read as a valid DOCX document") from exc

        parts: list[str] = [p.text for p in document.paragraphs if p.text.strip()]

        for table in document.tables:
            for row in table.rows:
                row_text = " | ".join(cell.text.strip() for cell in row.cells)
                if row_text.strip(" |"):
                    parts.append(row_text)

        return "\n".join(parts).strip()
